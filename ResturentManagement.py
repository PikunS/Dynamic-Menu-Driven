#--------------------hotel base program--------------------------

import mysql.connector as mq
import datetime
import docx		# pip install python-docx
import os

def myConnection():
	global mycon
	global mycur
	mycon=mq.connect(host="localhost", user="xxx", password="xxxxxxx")
	if mycon.is_connected():
		mycur=mycon.cursor()
		mycur.execute("create database if not exists management")
		mycur.execute("use management")
		mycur.execute("create table if not exists items(itno int(5) primary key, item_name varchar(40), rate decimal(6,2), foodtype varchar(20))")
		mycur.execute("create table if not exists Order_tb(itno int(5),qty int(4), rt decimal(6,2), tableno int(2))")
		mycur.execute("create table if not exists sales(billno int(5) primary key auto_increment, total_amt decimal (8,2), d_of_sales date)")
		mycur.execute("create table if not exists billdetails(bitno int(5) primary key auto_increment, billno int(5), itno int(5), itmname varchar(40) , qty int(5), rate decimal(6,2), d_of_sale date)")
		return True
	else:
		return False
# ------------------------ ORDER RELATED FUNCTION ---------------------------------
def OrderfoodItem(itnm=''):
	global mycon
	global mycur
	print(f'Welcome to our {itnm} menu')
	print("Give Your Order...")
	tableno=input("Enter Table Number: ")
	mycur.execute(f"select * from items where foodtype='{itnm}'")
	recs=mycur.fetchall()
	if recs:
		ch=True
		while ch:
			sno=1
			srn=[]
			l=[]
			rs=[]
			for rec in recs:
				print("{}: {} - (Rs.{})".format(sno, rec[1], rec[2]))
				srn.append(sno)
				l.append(rec[0])
				rs.append(rec[2])
				sno+=1
			print(f'{sno}: Return to Order Menu')
			opt1=int(input('Order Your Item Number: '))
			if opt1==sno:
				return
			elif opt1 in srn:
				qty=int(input("Enter Number of Cup/Glass/Plates : "))
				mycur.execute(f"insert into Order_tb values({l[opt1-1]},{qty},{rs[opt1-1]},{tableno})")
				mycon.commit()
			else:
				print("Sorry Wrong Item Number ")
			chh=input(f"Do you want to add more Item from {itnm}: Y/N ")
			if chh.upper()!='Y':
				ch=False
	else:
		print("No Item(s) Found ")

def orderMenu():
	global mycon
	global mycur
	while True:
		print('\033[1;30;41m____________________DOWN TOWN CAFE____________________\033[0;0m')
		print("\033[1;32;47m\t\t---- FOOD ORDER MENU -----\033[0;0m")
		mycur.execute("select distinct foodtype from items order by foodtype")
		recs=mycur.fetchall()
		sno=1
		l=[]
		if recs:
			for rec in recs:
				print(f'{sno}: {rec[0]}')
				l.append(rec[0])
				sno+=1
		print(f'{sno}: Return to main Menu')
		opt1=int(input('Please Enter Your Choice: '))
		if opt1==sno:
			return
		elif opt1<sno:
			OrderfoodItem(l[opt1-1])
		else:
			print('Invaild Option')
# ------------------------ END OF ORDER RELATED FUNCTION ---------------------------------

# ------------------------ FOOD ITEMS RELATED FUNCTION ---------------------------------
def addFoodItem():
	global mycon
	global mycur
	ch=True
	while ch:
		itemno = int(input('Enter item number: '))
		itemname=input('Enter item name:')
		rate = float(input('Enter rate: '))
		foodtype = input('Enter food type: ')
		food_item = mycur.execute("insert into items values ({}, '{}', {}, '{}')".format(itemno, itemname, rate, foodtype))
		mycon.commit()
		print("One Record Saved....")
		cc=input("Want to Add More Item : ")
		if cc.upper()!='Y':
			ch=False
	
def modifyFoodItem():
	global mycon
	global mycur
	ch=True
	while ch:
		itemno = int(input('Enter item number which you want to Modify : '))
		mycur.execute("select * from items where itno={}".format(itemno))
		recs = mycur.fetchone()
		if(recs):
			print('Item Number: {}'.format(itemno))
			itemname=input('Enter New Item Name:')
			rate = float(input('Enter New Rate: '))
			foodtype = input('Enter Food Type: ')
			mycur.execute("update items set item_name='{}', rate={}, foodtype='{}' where itno={}".format(itemname, rate, foodtype, itemno))
			mycon.commit()
			print("One Record Modified...")
		else:
			print("Item Number Not found")
		cc=input("Want to Modify More Item : (Y/N) ")
		if cc.upper()!='Y':
			ch=False
				
def deleFoodItem():
	global mycon
	global mycur
	ch=True
	while ch:
		itemno = int(input('Enter item number which you want to Delete : '))
		mycur.execute("select * from items where itno={}".format(itemno))
		if mycur.found_rows()>0:
			recs = mycur.fetchone()
			print('Item Number: {} is going to delete'.format(itemno))
			mycur.execute("delete from items where itno={}".format(itemno))
			mycon.commit()
			print("One Record Deleted...")
		else:
			print("Item Number Not found")
		cc=input("Want to Delete More Item : ")
		if cc.upper()!='Y':
			ch=False

def dispFoodItem():
	global mycon
	global mycur
	mycur.execute("select * from items")
	recs = mycur.fetchall()
	for rec in recs:
		print('Item Number: {}'.format(rec[0]))
		print('Item Name: {}'.format(rec[1]))
		print('Item Rate: {}'.format(rec[2]))
		print('Item Type: {}'.format(rec[3]))
	
def itemMenu():
	while True:
		print('\033[1;30;41m\____________________DOWN TOWN CAFE____________________\033[0;0m')
		print("\033[1;32;47m\t\t---- FOOD ITEM MENU -----\033[0;0m")
		print("\033[1;32m\t1: Add Food Items")
		print("\t2: Modify Food Items")
		print("\t3: Remove Food Items")
		print("\t4: Display Food Items")
		print("\t5: Return to Main Menu\033[0;0m")
		opt=int(input ("Enter your Choice... :"))
		if opt==1:
			addFoodItem()
		elif opt==2:
			modifyFoodItem()
		elif opt==3:
			deleFoodItem()
		elif opt==4:
			dispFoodItem()
		elif opt==5:
			return
# ------------------------ END OF FOOD RELATED FUNCTION ---------------------------------		

# ------------------------ BILL RELATED FUNCTION ---------------------------------
'''	s String format—this is the default type for strings
	d Decimal Integer. This uses a comma as the number separator character.
	n Number. This is the same as d except that it uses the current locale setting to insert the appropriate number separator characters.
	e Exponent notation. Prints the number in scientific notation using the letter e to indicate the exponent. The default precision is 6.
	f Fixed-point notation. Displays the number as a fixed-point number. The default precision is 6.
	% Percentage. Multiplies the number by 100 and displays in fixed ('f') format, followed by a percent sign.
'''
def BillgenAndPrint():
	global mycon
	global mycur
	tableno=input("Enter Table Number: ")
	mycur.execute(f"Select * from Order_tb where tableno={tableno}")
	recs=mycur.fetchall()
	totamt=0
	print('-------------------------------------- BILL --------------------------------------')
	
	dt=datetime.datetime.now()
	print(f'Bill No:={str(dt.year)+"001":40s} Date:{dt}')
	if recs:
		print(f"{'Item No':10s}{'Item Name':30s}{'Quantity':15s}{'Rate':14s}{'Amount'}")
		sno=1
		for rec in recs:
			mycur.execute(f"select item_name from items where itno={rec[0]}")
			itemname=mycur.fetchone()
			print(f"{sno:6d} {itemname[0]:25s} {rec[1]:10d} {rec[2]:15.2f} {(rec[1]*rec[2]):15.2f}")
			totamt+=(rec[1]*rec[2])
			sno+=1
		print('----------------------------------------------------------------------------------')
		print(f"{'Total Amount'}{totamt:>63.2f}")
		print('----------------------------------------------------------------------------------')
		amtpaid=float(input("Enter Amout to be Paid : "))
		if amtpaid==totamt:
			mycur.execute(f'insert into sales(total_amt,d_of_sales) values({totamt},\'{str(dt.year)+"-"+str(dt.month)+"-"+str(dt.day)}\')')
			mycon.commit()
			id=mycur.lastrowid
			mycur.execute(f"Select * from Order_tb where tableno={tableno}")
			recs=mycur.fetchall()
			for rec in recs:
				mycur.execute(f"select item_name from items where itno={rec[0]}")
				itemname=mycur.fetchone()
				mycur.execute(f'insert into billdetails(billno,itno,itmname,qty,rate,d_of_sale) values({id},{rec[0]},\'{itemname[0]}\',{rec[1]},{rec[2]},\'{str(dt.year)+"-"+str(dt.month)+"-"+str(dt.day)}\')')
			mycur.execute(f"delete from Order_tb where tableno={tableno}")
			mycon.commit()
			RePrintBill(id)
		

def RePrintBill(id=0):
	global mycon
	global mycur
	doc1 = docx.Document()
	mycur.execute(f"select * from sales where billno={id}")
	rec=mycur.fetchone()
	doc1.add_heading('DOWN TOWN CAFE',0)
	doc1.add_paragraph(f'Bill No:{str(rec[0]):40s} Date:{str(rec[2])}')
	doc1.add_paragraph("-"*80)
	doc1.add_paragraph(f"{'Item No':10s}{'Item Name':30s}{'Quantity':15s}{'Rate':14s}{'Amount'}")
	
	mycur.execute(f"Select * from billdetails where billno={id}")
	recs = mycur.fetchall()
	if mycur:
		sno=1
		totamt=0
		for rec in recs:			
			doc1.add_paragraph(f"{sno:6d} {rec[3]:25s} {rec[4]:10d} {rec[5]:15.2f} {(rec[4]*rec[5]):15.2f}")
			totamt+=(rec[4]*rec[5])
			sno+=1
		doc1.add_paragraph('-'*80)
		doc1.add_paragraph(f"{'Total Amount'}{totamt:>63.2f}")
		doc1.add_paragraph('-'*80')
		doc1.add_page_break()
		doc1.save("Result.docx")
		print("file Successfully Generated")
		os.system("Result.docx")
	else:
		print("No Record Found")

def BillMenu():
	while True:
		print("1: Generate Bill And Print")
		print("2: Reprint Bill")
		print("3: Return to Main Menu")
		opt=int(input ("Enter your Choice... :"))
		if opt==1:
			BillgenAndPrint()
		elif opt==2:
			id=int(input("Enter Bill id: "))
			RePrintBill(id)
		elif opt==3:
			return
# ------------------------ END BILL RELATED FUNCTION ---------------------------------
# ------------------------ REPORT RELATED FUNCTION ---------------------------------
def dailyReport():
	global mycon
	global mycur
	dtv=input("Enter date to see the report in 'yyyy-mm-dd' ")
	#dt=datetime.datetime.date(dtv)
	mycur.execute(f"select * from sales where d_of_sales='{dtv}'")
	recs = mycur.fetchall()
	if recs:
		print(("-"*24)+' DOWN TOWN CAFE '+("-"*24))
		print(f'Date:{dtv}')
		print("-"*70)
		daytot=0
		print(f"{'Bill No':40s}{'Amount'}")
		for rec in recs:
			print(f"{rec[0]:6d}{'':30s}{rec[1]:10.2f}")
			daytot+=rec[1]
		print("-"*70)
		print(f"{'Total Amount'}{daytot:>34.2f}")
		print("-"*70)
	else:
		print("Sorry No Transaction done in this date")
		
def monthlyReport():
	global mycon
	global mycur
	dtv=input("Enter Month Number to see the report in 'mm' ")
	#dt=datetime.datetime.date(dtv)
	mycur.execute(f"select d_of_sales, sum(total_amt) from sales where month(d_of_sales)={int(dtv)} group by d_of_sales")
	recs = mycur.fetchall()
	if recs:
		print(("-"*24)+' DOWN TOWN CAFE '+("-"*24))
		print(f'Month:{dtv}')
		print("-"*70)
		daytot=0
		print(f"{'Date':40s}{'Amount'}")
		for rec in recs:
			print(f"{str(rec[0]):20s}{'':16s}{rec[1]:10.2f}")
			daytot+=rec[1]
		print("-"*70)
		print(f"{'Total Amount'}{daytot:>34.2f}")
		print("-"*70)
	else:
		print("Sorry No Transaction done in this Month")
def yearlyReport():
	global mycon
	global mycur
	dtv=input("Enter Year Value to see the report in 'yyyy' ")
	#dt=datetime.datetime.date(dtv)
	mycur.execute(f"select month(d_of_sales), sum(total_amt) from sales where year(d_of_sales)={int(dtv)} group by month(d_of_sales)")
	recs = mycur.fetchall()
	if recs:
		print(("-"*24)+' DOWN TOWN CAFE '+("-"*24))
		print(f'Year:{dtv}')
		print("-"*70)
		daytot=0
		print(f"{'Month':40s}{'Amount'}")
		for rec in recs:
			print(f"{str(rec[0]):20s}{'':16s}{rec[1]:10.2f}")
			daytot+=rec[1]
		print("-"*70)
		print(f"{'Total Amount'}{daytot:>34.2f}")
		print("-"*70)
	else:
		print("Sorry No Transaction done in this Year")

def incomeReportMenu():
	while True:
		print("1: Daily Report")
		print("2: Monthly Report")
		print("3: Yearly Report")
		print("4: Return to Main Menu")
		opt=int(input("Enter Your Choice : "))
		if opt==1:
			dailyReport()
		elif opt==2:
			monthlyReport()
		elif opt==3:
			yearlyReport()
		elif opt==4:
			return
		else:
			print("Wrong Choice..")
# ------------------------ END OF REPORT RELATED FUNCTION ---------------------------------
# ------------------------ MAIN MENU ---------------------------------
def MainMenu():
	while True:
		print('\033[1;30;41m____________________DOWN TOWN CAFE____________________\033[0;0m')
		print("\033[1;31;47m\t\t---- MAIN MENU -----\033[0;0m")
		print("\033[1;32m\t1: Food Items\033[0;0m")
		print("\033[1;36m\t2: Food Order\033[0;0m")
		print("\033[1;31m\t3: Bill Details\033[0;0m")
		print("\033[1;33m\t4: Income Report\033[0;0m")
		print("\033[1;37m\t5: Exit from the Program \033[0;0m")
		opt=int(input("Enter your choice : "))
		if opt==1:
			itemMenu()
		elif opt==2:
			orderMenu()
		elif opt==3:
			BillMenu()
		elif opt==4:
			incomeReportMenu()
		elif opt==5:
			exit()
		else:
			print("Wrong Choice ")

		
if __name__ == '__main__':
	if myConnection():
		MainMenu()  
	else:                        
		print("Contact to your Manager")
